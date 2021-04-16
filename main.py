import tkinter as tk
from tkinter import filedialog
from docx import Document
from docx.shared import Pt
import os
import tkinter.messagebox
# from skimage import io
import hashlib


window=tk.Tk()
window.title("电子签名")
window.geometry('300x300')

var=tk.StringVar()
I=tk.Label(window,text='文件签名器',font=('Arial',15),width=15 ,height=2)
I.pack()
path = tk.StringVar()
passward1= tk.StringVar()
passward2= tk.StringVar()
passward=tk.StringVar()
dirname, filename = os.path.split(os.path.abspath(__file__)) 


def auto(passward):
    global path,dirname,filename
    if(os.path.isfile(path.get())==False):
        tk.messagebox.showerror(title = '错误！',message='文件错误')

    m = hashlib.md5()
    m.update(passward.encode("GBK"))#.encode("utf8")
    pic_name=m.hexdigest()
    
    for file in os.listdir(dirname+'//pic//'):
        if(pic_name==os.path.splitext(file)[0]):
            
            file_name=os.path.splitext(os.path.basename(path.get()))[0]
            file_dir=os.path.dirname(path.get())
            # 读取文档
            doc = Document(path.get())
            p = doc.add_paragraph()
            r = p.add_run()
            r.add_text('                               ')
            r.add_picture(dirname+'//pic//'+pic_name+'.jpg',width=Pt(250),height=Pt(100))
            # 存储文档
            doc.save(file_dir+"//"+file_name+"_new.docx")

    tk.messagebox.showinfo(title = '完成！',message='签名完成，保存在：'+file_dir+'\\\\'+file_name+"_new.docx")
def selectPath():
    global path
    #选择文件path_接收文件地址
    path_ = tk.filedialog.askopenfilename()
    #通过replace函数替换绝对文件地址中的/来使文件可被程序读取 
    #注意：\\转义后为\，所以\\\\转义后为\\
    path_=path_.replace("/","\\\\")
    #path设置path_的值
    path.set(path_)

def window2():
    global passward
    top_level = tk.Toplevel()
    top_level.title("新窗口")
    top_level.geometry('{}x{}+{}+{}'.format(500, 200, 500, 500)) #改变窗口位置和大小
    # tk.Label(top_level, text="输入密码",font=('Arial',15),width=15 ,height=2).pack()
    # passwor=tk.Text(top_level,width=25,height=2).pack()
    
    #输入框，标记，按键
    tk.Label(top_level,text = "目标路径:").grid(row = 0, column = 0)
    #输入框绑定变量path
    tk.Entry(top_level, textvariable = path).grid(row = 0, column = 1)
    tk.Button(top_level, text = "路径选择", command = selectPath).grid(row = 0, column = 2)
     #输入框，标记，按键
    tk.Label(top_level,text = "输入签名密码：").grid(row = 2, column = 0)
    tk.Entry(top_level,textvariable=passward).grid(row = 2, column = 1)
    tk.Button(top_level, text = "签名",width=25,height=2 ,command = lambda: auto(passward.get())).grid(row = 4, column = 1)

def add_auto_window():
    global path
    add_auto_window = tk.Toplevel()
    add_auto_window.title("文件签名")
    add_auto_window.geometry('{}x{}+{}+{}'.format(500, 200, 500, 500)) #改变窗口位置和大小

    #输入框，标记，按键
    tk.Label(add_auto_window,text = "图片路径:").grid(row = 0, column = 0)
    #输入框绑定变量path
    tk.Entry(add_auto_window, textvariable = path).grid(row = 0, column = 1)
    tk.Button(add_auto_window, text = "路径选择", command = selectPath).grid(row = 0, column = 2)

    #输入框，标记，按键
    tk.Label(add_auto_window,text = "设置签名密码：").grid(row = 2, column = 0)
    tk.Entry(add_auto_window,textvariable=passward1).grid(row = 2, column = 1)
    tk.Label(add_auto_window,text = "确认签名密码：").grid(row = 3, column = 0)
    tk.Entry(add_auto_window,textvariable=passward2).grid(row = 3, column = 1)

    tk.Button(add_auto_window, text = "添加",width=25,height=2 
    ,command = lambda: add_auto(passward1.get(), passward2.get())).grid(row = 4, column = 1)

def save_pic(p1):
    global path,dirname,filename
    #读取
    m = hashlib.md5()
    m.update(p1.encode("GBK"))
    pic_name=m.hexdigest()

    with open(path.get(), "rb") as r:
        with open(dirname+'\\pic\\'+pic_name+'.jpg','wb') as w:
            for line in r.readlines():
                w.write(line)
  




def add_auto(p1,p2):
    global path
    if(p1==p2):
        save_pic(p1)
        tk.messagebox.showinfo(title = '完成！',message='添加完成，快去签名吧')
    else:
        tk.messagebox.showerror(title = '失败！',message='密码错误或图片错误')
b1=tk.Button(window,text="添加签名",width=15 ,height=2,command=add_auto_window)
b2=tk.Button(window,text="文件签名",width=15 ,height=2,command=window2)
b1.pack()
b2.pack()



window.mainloop()

